#!/usr/bin/env python3
"""
Create Word document with team verification data
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# All teams data
ALL_TEAMS = {
    "Premier League": [
        {"name": "Arsenal", "instagram": "arsenal", "twitter": "Arsenal", "tiktok": "arsenal"},
        {"name": "Manchester City", "instagram": "mancity", "twitter": "ManCity", "tiktok": "mancity"},
        {"name": "Aston Villa", "instagram": "avfcofficial", "twitter": "AVFCOfficial", "tiktok": "avfcofficial"},
        {"name": "Chelsea", "instagram": "chelseafc", "twitter": "ChelseaFC", "tiktok": "chelseafc"},
        {"name": "Liverpool", "instagram": "liverpoolfc", "twitter": "LFC", "tiktok": "liverpoolfc"},
        {"name": "Sunderland", "instagram": "sunderlandafc", "twitter": "SunderlandAFC", "tiktok": "sunderlandafc"},
        {"name": "Manchester United", "instagram": "manchesterunited", "twitter": "ManUtd", "tiktok": "manutd"},
        {"name": "Crystal Palace", "instagram": "cpfc", "twitter": "CPFC", "tiktok": "cpfc"},
        {"name": "Brighton", "instagram": "officialbhafc", "twitter": "OfficialBHAFC", "tiktok": "officialbhafc"},
        {"name": "Everton", "instagram": "everton", "twitter": "Everton", "tiktok": "everton"},
        {"name": "Newcastle United", "instagram": "nufc", "twitter": "NUFC", "tiktok": "nufc"},
        {"name": "Brentford", "instagram": "brentfordfc", "twitter": "BrentfordFC", "tiktok": "brentfordfc"},
        {"name": "Tottenham Hotspur", "instagram": "spursofficial", "twitter": "SpursOfficial", "tiktok": "spursofficial"},
        {"name": "AFC Bournemouth", "instagram": "afcbournemouth", "twitter": "afcbournemouth", "tiktok": "afcb"},
        {"name": "Fulham", "instagram": "fulhamfc", "twitter": "FulhamFC", "tiktok": "fulhamfc"},
        {"name": "Leeds United", "instagram": "leedsunited", "twitter": "LUFC", "tiktok": "leedsunited"},
        {"name": "Nottingham Forest", "instagram": "nottinghamforest", "twitter": "NFFC", "tiktok": "nffc"},
        {"name": "West Ham United", "instagram": "westham", "twitter": "WestHam", "tiktok": "westham"},
        {"name": "Burnley", "instagram": "burnleyofficial", "twitter": "BurnleyOfficial", "tiktok": "burnleyofficial"},
        {"name": "Wolves", "instagram": "wolves", "twitter": "Wolves", "tiktok": "wolves"},
    ],
    "La Liga": [
        {"name": "FC Barcelona", "instagram": "fcbarcelona", "twitter": "FCBarcelona", "tiktok": "fcbarcelona"},
        {"name": "Real Madrid", "instagram": "realmadrid", "twitter": "realmadrid", "tiktok": "realmadrid"},
        {"name": "Atlético Madrid", "instagram": "atleticodemadrid", "twitter": "Atleti", "tiktok": "atleticodemadrid"},
        {"name": "Villarreal CF", "instagram": "villarrealcf", "twitter": "VillarrealCF", "tiktok": "villarrealcf"},
        {"name": "RCD Espanyol", "instagram": "rcdespanyol", "twitter": "RCDEspanyol", "tiktok": "rcdespanyol"},
        {"name": "Real Betis", "instagram": "realbetisbalompie", "twitter": "RealBetis", "tiktok": "realbetisbalompie"},
        {"name": "RC Celta", "instagram": "rccelta", "twitter": "RCCelta", "tiktok": "rccelta"},
        {"name": "Athletic Club", "instagram": "athleticclub", "twitter": "AthleticClub", "tiktok": "athleticclub"},
        {"name": "Elche CF", "instagram": "elchecfoficial", "twitter": "elchecf", "tiktok": "elchecfoficial"},
        {"name": "Sevilla FC", "instagram": "sevillafc", "twitter": "SevillaFC", "tiktok": "sevillafc"},
        {"name": "Getafe CF", "instagram": "getafecf", "twitter": "GetafeCF", "tiktok": "getafecf"},
        {"name": "CA Osasuna", "instagram": "caosasuna", "twitter": "CAOsasuna", "tiktok": "caosasuna"},
        {"name": "RCD Mallorca", "instagram": "rcdmallorca", "twitter": "RCD_Mallorca", "tiktok": "rcdmallorca"},
        {"name": "Deportivo Alavés", "instagram": "deportivoalaves", "twitter": "Alaves", "tiktok": "deportivoalaves"},
        {"name": "Rayo Vallecano", "instagram": "rayovallecano", "twitter": "RayoVallecano", "tiktok": "rayovallecano"},
        {"name": "Real Sociedad", "instagram": "realsociedad", "twitter": "RealSociedad", "tiktok": "realsociedad"},
        {"name": "Valencia CF", "instagram": "valenciacf", "twitter": "valenciacf", "tiktok": "valenciacf"},
        {"name": "Girona FC", "instagram": "gironafc", "twitter": "GironaFC", "tiktok": "gironafc"},
        {"name": "Real Oviedo", "instagram": "realoviedo", "twitter": "RealOviedo", "tiktok": "realoviedo"},
        {"name": "Levante UD", "instagram": "levanteud", "twitter": "LevanteUD", "tiktok": "levanteud"},
    ],
    "Serie A": [
        {"name": "Inter Milan", "instagram": "inter", "twitter": "Inter", "tiktok": "inter"},
        {"name": "AC Milan", "instagram": "acmilan", "twitter": "acmilan", "tiktok": "acmilan"},
        {"name": "SSC Napoli", "instagram": "officialsscnapoli", "twitter": "sscnapoli", "tiktok": "officialsscnapoli"},
        {"name": "AS Roma", "instagram": "asroma", "twitter": "ASRomaEN", "tiktok": "asroma"},
        {"name": "Juventus", "instagram": "juventus", "twitter": "juventusfcen", "tiktok": "juventus"},
        {"name": "Bologna FC", "instagram": "bolognafc1909", "twitter": "BolognaFC1909en", "tiktok": "bolognafc1909"},
        {"name": "Como 1907", "instagram": "comocalcio", "twitter": "Como_1907", "tiktok": "comocalcio"},
        {"name": "SS Lazio", "instagram": "officialsslazio", "twitter": "OfficialSSLazio", "tiktok": "officialsslazio"},
        {"name": "Atalanta BC", "instagram": "atalantabc", "twitter": "Atalanta_BC", "tiktok": "atalantabc"},
        {"name": "US Sassuolo", "instagram": "sassuolocalcio", "twitter": "SassuoloUS", "tiktok": "sassuolocalcio"},
        {"name": "US Cremonese", "instagram": "uscremonese", "twitter": "USCremonese", "tiktok": "uscremonese"},
        {"name": "Udinese Calcio", "instagram": "udinesecalcio", "twitter": "Udinese_1896", "tiktok": "udinesecalcio"},
        {"name": "Torino FC", "instagram": "torinofc_1906", "twitter": "TorinoFC_1906", "tiktok": "torinofc1906"},
        {"name": "US Lecce", "instagram": "uslecce", "twitter": "OfficialUSLecce", "tiktok": "uslecce"},
        {"name": "Cagliari Calcio", "instagram": "cagliaricalcio", "twitter": "CagliariCalcio", "tiktok": "cagliaricalcio"},
        {"name": "Parma Calcio", "instagram": "parmacalcio1913", "twitter": "1913parmacalcio", "tiktok": "parmacalcio1913"},
        {"name": "Genoa CFC", "instagram": "genoacfc", "twitter": "GenoaCFC", "tiktok": "genoacfc"},
        {"name": "Hellas Verona", "instagram": "hellasveronafcofficial", "twitter": "HellasVeronaFC", "tiktok": "hellasveronafcofficial"},
        {"name": "Pisa SC", "instagram": "pisasportingclub", "twitter": "PisaSC", "tiktok": "pisasporting"},
        {"name": "ACF Fiorentina", "instagram": "acffiorentina", "twitter": "acffiorentina", "tiktok": "acffiorentina"},
    ],
    "Ligue 1": [
        {"name": "RC Lens", "instagram": "rclens", "twitter": "RCLens", "tiktok": "rclens"},
        {"name": "Paris Saint-Germain", "instagram": "psg", "twitter": "PSG_English", "tiktok": "psg"},
        {"name": "Olympique de Marseille", "instagram": "olympiquedemarseille", "twitter": "OM_English", "tiktok": "om"},
        {"name": "LOSC Lille", "instagram": "loscofficiel", "twitter": "LOSC_EN", "tiktok": "loscofficiel"},
        {"name": "Olympique Lyonnais", "instagram": "ol", "twitter": "OL", "tiktok": "ol"},
        {"name": "Stade Rennais", "instagram": "staderennaisfc", "twitter": "staderennais", "tiktok": "staderennaisfc"},
        {"name": "RC Strasbourg", "instagram": "rcsa", "twitter": "RCSA", "tiktok": "rcstrasbourgalsace"},
        {"name": "Toulouse FC", "instagram": "toulousefc", "twitter": "ToulouseFC", "tiktok": "toulousefc"},
        {"name": "AS Monaco", "instagram": "asmonaco", "twitter": "AS_Monaco_EN", "tiktok": "asmonaco"},
        {"name": "Angers SCO", "instagram": "angersscofficiel", "twitter": "AngersSCO", "tiktok": "angersscofficiel"},
        {"name": "Stade Brestois", "instagram": "sb29", "twitter": "SB29", "tiktok": "sb29official"},
        {"name": "FC Lorient", "instagram": "fclorient", "twitter": "FCLorient", "tiktok": "fclorient"},
        {"name": "OGC Nice", "instagram": "ogcnice", "twitter": "ogcnice", "tiktok": "ogcnice"},
        {"name": "Paris FC", "instagram": "parisfcofficial", "twitter": "ParisFC", "tiktok": "parisfc"},
        {"name": "Le Havre AC", "instagram": "hac_officiel", "twitter": "HAC_Foot", "tiktok": "lehavreac"},
        {"name": "AJ Auxerre", "instagram": "ajauxerre", "twitter": "AJA", "tiktok": "ajauxerre"},
        {"name": "FC Nantes", "instagram": "fcnantes", "twitter": "FCNantes", "tiktok": "fcnantes"},
        {"name": "FC Metz", "instagram": "fcmetz", "twitter": "FCMetz", "tiktok": "fcmetz"},
    ],
    "Süper Lig": [
        {"name": "Galatasaray", "instagram": "galatasaray", "twitter": "GalatasaraySK", "tiktok": "galatasaray"},
        {"name": "Fenerbahçe", "instagram": "fenerbahce", "twitter": "Fenerbahce", "tiktok": "fenerbahce"},
        {"name": "Trabzonspor", "instagram": "trabzonspor", "twitter": "Trabzonspor", "tiktok": "trabzonspor"},
        {"name": "Göztepe", "instagram": "goztepeskresmi", "twitter": "Goztepe", "tiktok": "goztepeskresmi"},
        {"name": "Beşiktaş", "instagram": "besiktas", "twitter": "Besiktas", "tiktok": "bjk"},
        {"name": "Samsunspor", "instagram": "samsunspor", "twitter": "SamsunSporKlb", "tiktok": "samsunspor"},
        {"name": "İstanbul Başakşehir", "instagram": "ibfk2014", "twitter": "ibfk2014", "tiktok": "basaksehir"},
        {"name": "Kocaelispor", "instagram": "kocaelisporkulubu", "twitter": "kocaelispor", "tiktok": "kocaelisporkulubu"},
        {"name": "Gaziantep FK", "instagram": "gaziantepfk", "twitter": "GaziantepFK", "tiktok": "gaziantepfk"},
        {"name": "Alanyaspor", "instagram": "alanyaspor", "twitter": "Alanyaspor", "tiktok": "alanyaspor"},
        {"name": "Gençlerbirliği", "instagram": "genclerbirligi", "twitter": "Genclerbirligi", "tiktok": "genclerbirligiresmi"},
        {"name": "Çaykur Rizespor", "instagram": "caykurrizespor", "twitter": "CRizespor", "tiktok": "caykurrizespor"},
        {"name": "Konyaspor", "instagram": "konyaspor", "twitter": "konyaspor", "tiktok": "konyaspor"},
        {"name": "Kasımpaşa", "instagram": "kasimpasask", "twitter": "Kasimpasa", "tiktok": "kasimpasa"},
        {"name": "Antalyaspor", "instagram": "antalyaspor", "twitter": "Antalyaspor", "tiktok": "antalyaspor"},
        {"name": "Kayserispor", "instagram": "kayserispor", "twitter": "Kayserispor", "tiktok": "kayserispor"},
        {"name": "Eyüpspor", "instagram": "eyupspor", "twitter": "Eyupspor", "tiktok": "eyupspor1925"},
        {"name": "Fatih Karagümrük", "instagram": "karagumruksk", "twitter": "karagumruksk", "tiktok": "karagumruksk"},
    ],
}

def create_word_document():
    doc = Document()
    
    # Title
    title = doc.add_heading('Team Social Media Verification Data', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0].font
    subtitle_format.size = Pt(10)
    subtitle_format.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()
    
    # Add data for each league
    for league, teams in ALL_TEAMS.items():
        # League heading
        doc.add_heading(f'{league} ({len(teams)} teams)', 1)
        
        # Create table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Team'
        header_cells[1].text = 'Instagram'
        header_cells[2].text = 'Twitter/X'
        header_cells[3].text = 'TikTok'
        
        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Add team data
        for team in teams:
            row_cells = table.add_row().cells
            row_cells[0].text = team['name']
            row_cells[1].text = f"@{team['instagram']}"
            row_cells[2].text = f"@{team['twitter']}"
            row_cells[3].text = f"@{team['tiktok']}"
        
        doc.add_paragraph()
    
    # Save document
    filename = 'team_verification_data.docx'
    doc.save(filename)
    print(f"✅ Word document created: {filename}")
    
    # Summary
    total_teams = sum(len(teams) for teams in ALL_TEAMS.values())
    print(f"\n📊 Summary:")
    print(f"   Total teams: {total_teams}")
    print(f"   Total leagues: {len(ALL_TEAMS)}")

if __name__ == "__main__":
    create_word_document()
